"""Generate the course report DOCX.

Output naming follows: 姓名_学号_报告名称.docx
"""
from __future__ import annotations

import json
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
PROJ = ROOT.parent
GALLERY = PROJ / "结果" / "gallery"
OUT = PROJ / "滕向东_10235101422_家庭老照片修复与纪念化生成课程报告.docx"


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=None, bold=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph_font(paragraph, east_asia="宋体", ascii_font="Times New Roman", size=10.5):
    for run in paragraph.runs:
        set_run_font(run, east_asia, ascii_font, size)


def set_cell_text(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, "黑体", "Arial", size=16, bold=True)
    elif level == 2:
        set_run_font(run, "黑体", "Arial", size=14, bold=True)
    else:
        set_run_font(run, "黑体", "Arial", size=12, bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-12)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("• " + text)
    set_run_font(run, size=10.5)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=10)
        shade_cell(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=9.5)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_image(doc, filename, caption, width_cm=14.5):
    path = GALLERY / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, "宋体", "Times New Roman", size=9)


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("数字图像处理课程项目报告")
    set_run_font(r, "黑体", "Arial", size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("家庭老照片修复与纪念化生成")
    set_run_font(r, "黑体", "Arial", size=18, bold=True)

    for _ in range(5):
        doc.add_paragraph()

    info = [
        ("课程", "数字图像处理_2026"),
        ("学院", "软件工程学院"),
        ("姓名", "滕向东"),
        ("学号", "10235101422"),
        ("指导教师", "曹桂涛 教授"),
        ("日期", "2026 年 6 月"),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (k, v) in zip(table.rows, info):
        set_cell_text(row.cells[0], k, bold=True, size=12)
        set_cell_text(row.cells[1], v, size=12)
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(8)
    doc.add_page_break()


def build_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.3)
    sec.left_margin = Cm(2.7)
    sec.right_margin = Cm(2.7)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    add_cover(doc)

    add_heading(doc, "一、项目概述", 1)
    add_body(doc, "本项目实现了一套面向家庭老照片修复与纪念化生成的桌面应用。用户导入旧照后，可以根据照片特点选择修复模式，改善偏黄、褪色、灰蒙、颗粒、模糊、折痕划痕等问题，并生成相册、胶片、明信片、淡彩手绘等纪念版图像。")
    add_body(doc, "项目目标可以概括为：修复可用、输出可展示、算法可解释。系统既面向真实照片处理需求，也保留了数字图像处理课程中常见算法的可展示性。")

    add_table(doc, ["模块", "面向需求", "主要功能", "输出"], [
        ["家庭老照片修复", "照片质量退化", "温和修复、黄化修复、黑白增强、颗粒修复、人像清晰化、划痕修复", "修复图"],
        ["纪念版生成", "家庭展示与保存", "复古相册、暖色胶片、纪念明信片、淡彩手绘", "纪念版图像"],
        ["VGG 自定义艺术风格", "高级个性化", "Gatys 风格迁移、任意风格图", "艺术风格图"],
        ["高级算法功能", "实验与调参", "增强、去噪、锐化、恢复、频域、小波等", "单算法结果"],
    ], widths=[3, 3, 7, 3])

    add_heading(doc, "二、需求分析", 1)
    add_body(doc, "家庭老照片具有情感价值，因此修复时不能只追求强烈增强。系统需要改善画质，同时保留旧照质感。常见需求如下表。")
    add_table(doc, ["需求", "用户看到的问题", "系统模式", "设计目标"], [
        ["通用改善", "轻微发黄、灰蒙、颗粒", "温和修复", "稳定自然，不过度处理"],
        ["色彩老化", "偏黄、偏棕、褪色", "泛黄褪色修复", "去除过量黄化，保留暖色"],
        ["黑白旧照", "灰度层次弱、边缘不清", "黑白旧照增强", "提高明暗层次和相纸质感"],
        ["扫描噪声", "背景颗粒、暗部脏点", "扫描颗粒修复", "去噪并保护人物轮廓"],
        ["人像不清", "脸部和轮廓略模糊", "人像清晰化", "增强中等边缘，抑制噪声放大"],
        ["局部破损", "折痕、划痕、白线", "折痕划痕辅助修复", "自动辅助掩膜 + 用户掩膜 + inpaint"],
    ], widths=[2.6, 4, 3.2, 5])

    add_heading(doc, "三、核心模块设计", 1)
    add_heading(doc, "3.1 家庭老照片修复模块", 2)
    add_body(doc, "该模块位于 family_restore.py。每个函数都不是单一算法，而是围绕一种家庭旧照问题组合出固定处理流程。")
    add_table(doc, ["模式", "处理流程", "技术目的"], [
        ["温和修复", "轻度 NLM → 黄化校正 → CLAHE → 亮度 USM", "给多数旧照提供自然初始结果"],
        ["泛黄褪色修复", "LAB 去黄 → L 通道增强 → 饱和恢复", "校正黄化并保留少量暖色"],
        ["黑白旧照增强", "灰度去噪 → CLAHE → S 曲线 → 暖纸色", "增强层次和旧相纸观感"],
        ["扫描颗粒修复", "YCrCb 分离 → 色度强去噪 → 亮度弱去噪 → 边缘融合", "减少颗粒并保留人物边缘"],
        ["人像清晰化", "双边基础层 → 细节层 → 梯度 mask → 亮度 gate", "提升轮廓，避免暗部噪声锐化"],
        ["折痕划痕辅助修复", "top-hat / black-hat → 掩膜合并 → inpaint", "修补细线划痕和局部破损"],
    ], widths=[3, 6.5, 6])
    add_image(doc, "custom_温和修复.jpg", "图 1  温和修复前后对比")
    add_image(doc, "custom_泛黄褪色修复.jpg", "图 2  泛黄褪色修复前后对比")
    add_image(doc, "custom_扫描颗粒修复.jpg", "图 3  扫描颗粒修复前后对比")

    add_heading(doc, "3.2 纪念版生成模块", 2)
    add_body(doc, "该模块位于 commemorative.py，使用传统图像处理生成家庭照片纪念模板。模板目标是稳定、快速、贴合家庭相册和纪念展示场景。")
    add_table(doc, ["模板", "处理方式", "适用场景"], [
        ["复古相册风", "低饱和、暖色曲线、暗角、米色边框", "家庭相册保存"],
        ["暖色胶片风", "S 曲线、暖高光、轻微颗粒、暗角", "怀旧生活照"],
        ["纪念明信片风", "米色画布、白边、投影、标题区域", "展示与礼物卡片"],
        ["淡彩手绘风", "双边平滑、浅线稿、色彩融合", "柔和纪念图"],
    ], widths=[3, 7, 5])
    add_image(doc, "memory_复古相册风.jpg", "图 4  复古相册风纪念版")
    add_image(doc, "memory_纪念明信片风.jpg", "图 5  纪念明信片风纪念版")

    add_heading(doc, "3.3 VGG 自定义艺术风格模块", 2)
    add_body(doc, "该模块位于 neural_style.py。系统使用预训练 VGG19 提取内容特征和风格特征，不训练网络。内容由深层特征图表示，风格由多层 Gram 矩阵表示，最终通过 L-BFGS 优化输出图像素。")
    add_bullet(doc, "输入：当前修复图 + 用户选择的风格图。")
    add_bullet(doc, "输出：保留照片结构、具有参考图纹理和色彩倾向的艺术风格图。")
    add_bullet(doc, "工程处理：后台线程运行，避免 tkinter 主界面卡死；处理尺寸默认 512px，可调至 720px。")
    add_image(doc, "vgg_家庭暖调风格迁移.jpg", "图 6  VGG 自定义艺术风格迁移示例")

    add_heading(doc, "3.4 高级算法功能", 2)
    add_body(doc, "高级算法功能用于课程展示、参数实验和进一步精修。它包括基础操作、对比度增强、平滑去噪、锐化、图像恢复、频域处理、小波去噪、分割和形态学等。")
    add_table(doc, ["类别", "功能示例", "作用"], [
        ["增强", "CLAHE、Gamma、白平衡", "改善灰蒙、偏暗和偏色"],
        ["去噪", "中值、双边、NLM", "抑制颗粒和扫描噪声"],
        ["锐化", "USM、Laplacian、Sobel", "增强轮廓和边缘"],
        ["恢复", "维纳滤波、inpaint", "处理模糊和局部破损"],
        ["频域/小波", "傅里叶频谱、低通、高通、小波去噪", "从变换域理解图像退化"],
    ], widths=[3, 6, 6])

    add_heading(doc, "四、系统实现与界面", 1)
    add_body(doc, "系统使用 tkinter 实现桌面界面，左侧按使用顺序组织为家庭老照片修复、纪念版生成和高级算法功能。中间为原图与当前结果并排预览，下方为动态参数区。")
    add_table(doc, ["界面区域", "目的", "说明"], [
        ["家庭老照片修复", "快速修复", "选择定制修复模式，直接得到处理结果"],
        ["纪念版生成", "展示输出", "将当前修复图转换为纪念模板或 VGG 风格图"],
        ["高级算法功能", "调参实验", "展示单个传统图像处理算法效果"],
        ["预览区", "效果对比", "原图和当前结果并排显示"],
        ["导出功能", "保存成果", "保存处理结果或导出前后对比图"],
    ], widths=[3.2, 3, 8])

    add_heading(doc, "五、实验结果", 1)
    add_body(doc, "实验采用合成老照片和真实老照片双轨评估。合成样张有干净原图，可以计算 PSNR 和 SSIM；真实照片没有标准答案，主要观察偏色、噪声、清晰度、自然度和纪念效果。")
    metrics_path = GALLERY / "metrics.json"
    if metrics_path.exists():
        metrics = json.loads(metrics_path.read_text(encoding="utf-8"))
        rows = [[m["方法"], f'{m["PSNR (dB)"]:.2f}', f'{m["SSIM"]:.4f}'] for m in metrics]
        add_table(doc, ["方法", "PSNR (dB)", "SSIM"], rows, widths=[7, 4, 4])
    add_image(doc, "metrics_chart.jpg", "图 7  合成样张 PSNR / SSIM 指标对比")
    add_body(doc, "指标结果用于说明算法在合成退化样张上的表现。对于真实老照片，色彩校正、锐化和纪念模板会主动改变像素，因此不能只用 PSNR 判断效果。")

    add_heading(doc, "六、运行方式", 1)
    add_body(doc, "项目运行环境为 conda ai-service。推荐使用绝对路径启动，避免当前目录不同导致找不到 main.py。")
    add_bullet(doc, r"启动 App：E:\miniconda3\envs\ai-service\python.exe E:\数字图像处理\项目\代码\main.py")
    add_bullet(doc, r"重新生成展示图：E:\miniconda3\envs\ai-service\python.exe E:\数字图像处理\项目\代码\generate_gallery.py")
    add_bullet(doc, r"重新生成幻灯片：E:\miniconda3\envs\ai-service\python.exe E:\数字图像处理\项目\代码\build_slides.py")

    add_heading(doc, "七、总结", 1)
    add_body(doc, "本项目完成了家庭老照片修复与纪念化生成系统。家庭老照片修复模块解决照片质量问题，纪念版生成模块解决展示和保存需求，VGG 自定义艺术风格模块提供高级个性化扩展，高级算法功能用于课程展示和进一步调参。")
    add_body(doc, "系统达到了修复可用、输出可展示、算法可解释的目标，体现了数字图像处理方法在实际照片修复场景中的应用价值。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
